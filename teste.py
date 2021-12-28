from docx import Document
from docx.shared import Inches

class DocumentBuilder():
    document = Document()
    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()
    document.save('demo.docx')


class formateText():
    def __init__(self, text) -> None:
        self.text = text

    command_flags = dict(
        negrito = False, 
        itálico = False, 
        sublinhado = False
    )

    command_action = dict(
        negrito = ["[b]", "[/b]"],
        itálico = ["[i]", "[/i]"],
        sublinhado = ["[u]", "[/u]"],
        vírgula = ",",
        interrogação = "?"
    )

    def formate(self):
        commandKeys = self.command_flags.keys()
        limit = len(self.text) - 1

        for i, str in enumerate(self.text):
            if "ativar" == str.lower() and i + 1 < limit:
                command = self.text[i + 1].lower()
                if command in commandKeys and not self.command_flags[command]:
                    del(self.text[i])
                    self.text[i] = self.command_action.get(command)[0]
                    self.command_flags[command] = True

            elif "desativar" == str.lower() and i + 1 < limit:
                command = self.text[i + 1].lower()
                if command in commandKeys and self.command_flags[command]:
                    del(self.text[i])
                    self.text[i] = self.command_action.get(command)[1]
                    self.command_flags[command] = False

            elif "vírgula" in str:
                del(self.text[i])
                self.text[i - 1] += self.command_action[str]

            elif "interrogação" in str:
                del(self.text[i])
                self.text[i - 1] += self.command_action[str]

        return(" ".join(self.text))
    
text = ["ativar", "negrito", "testando", "texto", "em", "negrito", "desativar", "negrito"]
a = formateText(text).formate()
print(a)

b = DocumentBuilder()
