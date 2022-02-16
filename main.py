from kivy.app import App
from pathlib import Path
from typing import Union
from docx import Document
from kivy.metrics import sp
from kivy.clock import Clock
from kivy.config import Config
import speech_recognition as sr
from kivy.uix.popup import Popup
from kivy.uix.label import Label
from kivy.uix.image import Image
from kivy.uix.switch import Switch
from kivy.core.window import Window
from kivy.animation import Animation
import os, sys, json, random, platform
from kivy.uix.textinput import TextInput
from kivy.uix.boxlayout import BoxLayout
from kivy.properties import ListProperty
from kivy.graphics import RoundedRectangle, Color
from kivy.uix.behaviors.button import ButtonBehavior
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.resources import resource_add_path, resource_find

Config.set('kivy', 'window_icon', 'icons/logo.png')
Config.write()

r = sr.Recognizer()
m = sr.Microphone()


class Manager(ScreenManager):
    pass


class DocumentBuilder():
    def __init__(self) -> None:
        self.doc = Document()

    def fileTitle(self, newTitle) -> None:
        self.FileTitle = newTitle

    def docTitle(self, newTitle) -> None:
        self.doc.add_heading(newTitle, 0)

    def newParagraph(self, textList) -> None:
        command_flags = dict(
            negrito=False,
            itálico=False,
            sublinhado=False
        )
        p = self.doc.add_paragraph()
        wordsList = []

        for i in textList:
            wordsList.extend(i.split())

        for str in wordsList:
            if len(str) > 3:
                if "[b]" in str:
                    command_flags["negrito"] = True
                    str = str.split("[b]")[1]
                if "[i]" in str:
                    command_flags["itálico"] = True
                    str = str.split("[i]")[1]
                if "[u]" in str:
                    command_flags["sublinhado"] = True
                    str = str.split("[u]")[1]

                if "[/b]" in str:
                    command_flags["negrito"] = False
                    p.add_run(str.split("[/b]")[0]).bold = True
                    str = ""
                if "[/i]" in str:
                    command_flags["itálico"] = False
                    p.add_run(str.split("[/i]")[0]).italic = True
                    str = ""
                if "[/u]" in str:
                    command_flags["sublinhado"] = False
                    p.add_run(str.split("[/u]")[0]).underline = True
                    str = ""

            runner = p.add_run(str + " ")
            if command_flags["negrito"] == True:
                runner.bold = True
            if command_flags["itálico"] == True:
                runner.italic = True
            if command_flags["sublinhado"] == True:
                runner.underline = True

    def builder(self) -> None:
        self.doc.add_page_break()
        path, home = "", ""
        so = platform.system()

        if so == "Windows":
            home = Path.home() / "Documents"
        elif so == "Linux":
            home = Path.home() / "Documentos"

        path = os.path.join(home, "Yumi")

        try:
            os.mkdir(path)
        except OSError as error:
            print(error)

        self.doc.save("{}/{}.docx".format(path, self.FileTitle))


class FormateText():
    def __init__(self, text) -> None:
        self.text = text

    command_flags = dict(
        negrito=False,
        itálico=False,
        sublinhado=False
    )

    command_action = dict(
        negrito=["[b]", "[/b]"],
        itálico=["[i]", "[/i]"],
        sublinhado=["[u]", "[/u]"],
        vírgula=",",
        interrogação="?"
    )

    def formate(self) -> Union[str, bool]:
        commandKeys = self.command_flags.keys()
        limit = len(self.text) - 1
        newParagraph = False
        delParagraph = False
        correct = False

        for i, str in enumerate(self.text):
            if "vírgula" in str.lower():
                del(self.text[i])
                self.text[i - 1] += self.command_action[str]
            elif "interrogação" in str.lower():
                del(self.text[i])
                self.text[i - 1] += self.command_action[str]
            elif "novo" in str.lower() and i + 1 < limit:
                command = self.text[i + 1].lower()
                if command == "parágrafo":
                    del(self.text[i + 1])
                    del(self.text[i])
                    newParagraph = True
            elif "ativar" == str.lower() and i + 1 < limit:
                command = self.text[i + 1].lower()
                if command in commandKeys and not self.command_flags[command]:
                    del(self.text[i + 1])
                    del(self.text[i])
                    self.text[i] = self.command_action[command][0] + \
                        self.text[i]
                    self.command_flags[command] = True
            elif "desativar" == str.lower() and i + 1 < limit:
                command = self.text[i + 1].lower()
                if command in commandKeys and self.command_flags[command]:
                    del(self.text[i + 1])
                    del(self.text[i])
                    self.text[i - 1] += self.command_action[command][1]
                    self.command_flags[command] = False
            elif "apagar" in str.lower() and i + 1 <= limit:
                command = self.text[i + 1].lower()
                if command == "parágrafo":
                    delParagraph = True
            elif "corrigir" in str.lower() and i + 1 <= limit:
                if(self.text[i + 1].lower() == "por"):
                    del(self.text[i + 1])
                    del(self.text[i])
                    correct = True

        return([" ".join(self.text), newParagraph, delParagraph, correct])


class Menu(Screen):
    def on_pre_enter(self) -> None:
        Window.bind(on_request_close=self.confirm_exit)

    def confirm_exit(self, *args, **kwargs) -> bool:
        box = BoxLayout(orientation="vertical", padding=10, spacing=10)
        popup = Popup(
            title="Deseja realmente sair?",
            content=box,
            size_hint=(None, None),
            size=(sp(150), sp(100))
        )
        buttons = BoxLayout(padding=sp(5), spacing=sp(10))
        exitButton = RoundButton(
            text="Sim", on_release=App.get_running_app().stop)
        closeButton = RoundButton(text="Não", on_release=popup.dismiss)

        buttons.add_widget(exitButton)
        buttons.add_widget(closeButton)
        box.add_widget(Image(source="icons/warning.png"))
        box.add_widget(buttons)

        animation = Animation(size=(sp(300), sp(200)),
                              duration=0.3, t="out_back")
        animation.start(popup)
        popup.open()

        return True


class RoundButton(ButtonBehavior, Label):
    cor = ListProperty([0.1, 0.5, 0.8, 1])
    cor2 = ListProperty([0.3, 0.1, 0.9, 1])

    def __init__(self, **kwargs) -> None:
        super(RoundButton, self).__init__(**kwargs)
        self.update()

    def on_pos(self, *args) -> None:
        self.update()

    def on_size(self, *args) -> None:
        self.update()

    def on_press(self, *args) -> None:
        self.cor, self.cor2 = self.cor2, self.cor

    def on_release(self, *args) -> None:
        self.cor, self.cor2 = self.cor2, self.cor

    def on_cor(self, *args) -> None:
        self.update()

    def update(self, *args) -> None:
        self.canvas.before.clear()
        with self.canvas.before:
            Color(rgba=self.cor)
            RoundedRectangle(pos=self.pos, size=self.size, radius=[20])


class TextBoxContainer(Screen):
    pheases = [[]]
    path = ''

    def on_pre_enter(self, *args) -> None:
        super().on_pre_enter(*args)
        self.ids.textBox.clear_widgets()
        self.path = App.get_running_app().user_data_dir + "/"
        self.readData()
        Window.bind(on_keyboard=self.keyInterrupt)
        for phease in self.pheases:
            self.ids.textBox.add_widget(LabelBox(text=" ".join(phease)))

    def keyInterrupt(self, window, key, *args) -> bool:
        if key == 32:
            self.initRecorder()
        if key == 27:
            App.get_running_app().root.current = "menu"
        
        return True

    def on_pre_leave(self, *args) -> None:
        super().on_pre_leave(*args)
        Window.unbind(on_keyboard=self.keyInterrupt)

    def readData(self, *args) -> None:
        try:
            with open(self.path + "data.json", "r") as data:
                self.pheases = json.load(data)
        except FileNotFoundError:
            print("Arquivo não encontrado no caminho: {}".format(self.path))

    def saveData(self) -> None:
        with open(self.path + "data.json", "w") as data:
            json.dump(self.pheases, data)

    def docxBuilder(self, *args) -> None:
        hash = random.getrandbits(32)
        title = self.userInput.text or "doc_%08x" % hash
        document = DocumentBuilder()
        document.fileTitle(title)

        if self.switchValue:
            document.docTitle(title)

        for textList in self.pheases:
            document.newParagraph(textList)

        document.builder()
        self.popup.dismiss()
        Clock.schedule_once(self.createdDocPopup, 0.5)

    def createdDocPopup(self, *args) -> bool:
        so = platform.system()
        home = ""

        if so == "Windows":
            home = Path.home() / "Documents"
        elif so == "Linux":
            home = Path.home() / "Documentos"

        box = BoxLayout(orientation="horizontal", padding=10, spacing=10)
        popup = Popup(
            title="Sucesso!",
            content=box,
            size_hint=(None, None),
            size=(sp(350), sp(100)),
            pos_hint={"top": 0.9}
        )

        labelBox = BoxLayout(orientation="vertical")
        label1 = Label(text="Documento salvo em:", bold=True, font_size='16sp')
        label2 = Label(text=str(home) + '/Yumi')
        labelBox.add_widget(label1)
        labelBox.add_widget(label2)
        box.add_widget(Image(source="icons/logo_com_nome.png"))
        box.add_widget(labelBox)
        animation = Animation(size=(sp(530), sp(150)),
                              duration=0.3, t="out_back")
        animation.start(popup)

        popup.open()
        return True

    def docPopup(self, *args, **kwargs) -> bool:
        self.switchValue = True

        def updateSwitchValue(value) -> None:
            self.switchValue = value

        def onChange(self, value) -> None:
            updateSwitchValue(True if value else False)

        box = BoxLayout(orientation="vertical", padding=10, spacing=10)
        self.popup = Popup(
            title="Nome do documento",
            content=box,
            size_hint=(None, None),
            size=(sp(200), sp(200)))
        self.userInput = UserInput()
        buttons = BoxLayout(padding=sp(5), spacing=sp(10))
        box2 = BoxLayout(spacing=sp(10))
        label = Label(text="Usar como Título")
        switch = Switch(active=True)
        switch.bind(active=onChange)
        exitButton = RoundButton(text="Exportar", on_release=self.docxBuilder)

        buttons.add_widget(exitButton)
        box2.add_widget(label)
        box2.add_widget(switch)
        box.add_widget(self.userInput)
        box.add_widget(box2)
        box.add_widget(buttons)

        animation = Animation(size=(sp(320), sp(220)),
                              duration=0.3, t="out_back")
        animation.start(self.popup)
        self.popup.open()

        return True

    def messageError(self, msg, *args) -> bool:
        self.ids.image_mic.source = "icons/mic.png"
        box = BoxLayout(orientation="vertical", padding=10, spacing=10)
        popup = Popup(
            title="Ops, ocorreu um erro!",
            content=box,
            size_hint=(None, None),
            size=(sp(150), sp(150))
        )

        labelMsg = Label(text=msg, bold=True, font_size='16sp')
        box.add_widget(Image(source="icons/logo.png"))
        box.add_widget(labelMsg)
        animation = Animation(size=(sp(330), sp(200)),
                              duration=0.3, t="out_back")
        animation.start(popup)

        popup.open()
        return True

    def message(self, msg, *args) -> None:
        [phease, newParagraph, delParagraph, correct] = FormateText(
            msg.split()).formate()
        self.ids.image_mic.source = "icons/mic.png"
        textBoxTree = self.ids.textBox.children
        pheasesLen = len(self.pheases)

        if newParagraph:
            self.ids.textBox.add_widget(LabelBox(text=phease))
            self.pheases.append([phease])
        elif delParagraph:
            self.pheases.pop()
            self.ids.textBox.remove_widget(self.ids.textBox.children[0])
        elif correct:
            list_index = pheasesLen - 1
            self.pheases[list_index].append(phease)
            aux_list = self.pheases[list_index]
            aux_list.pop(len(aux_list) - 2)
            self.pheases[list_index] = aux_list
            self.ids.textBox.children[0].text = " ".join(aux_list)
        elif len(textBoxTree) > 0:
            self.ids.textBox.children[0].text += " " + phease
            self.pheases[pheasesLen - 1].append(phease)
        else:
            self.ids.textBox.add_widget(LabelBox(text=phease))
            self.pheases[pheasesLen - 1].append(phease)

        # self.saveData()

        # print("Lista de frases:\n",self.pheases)
        # print("frase antes das mudanças:\n", msg)
        # print("frase pré-processada:\n", phease)

    def listen(self, *args) -> None:
        with m as source:
            audio = r.listen(source)

        try:
            # Reconhecendo com API do google
            value = r.recognize_google(audio, language='pt-BR')
            self.message(value)

        except sr.UnknownValueError:
            self.messageError("Não entendi o que você falou.")

        except sr.RequestError as e:
            self.messageError("Não consigo conectar ao servidor :(")
            print("ERROR ====> {0}".format(e))

    def initRecorder(self, *args) -> None:
        self.ids.image_mic.source = "icons/mic_active.png"
        # Atrasa a execução da método listen() em 0.3 segundos
        Clock.schedule_once(self.listen, 0.3)


class Instructions(Screen):
    def on_pre_enter(self, *args) -> None:
        super().on_pre_enter(*args)
        self.ids.textBoxIntructions.clear_widgets()

        title = "Introdução"
        subTitle = "Comandos"

        intro = ("Para que o reconhecimento de fala funcione"
        " corretamente, é muito importante que você esteja em um lugar"
        " com pouco barulho. Além disso, verifique nas configurações"
        " de seu microfone se o volume está muito alto (volume muito"
        " alto pode ser prejudicial, pois faz o microfone capturar muitos"
        " ruídos e portanto, dificulta o sistema de reconhecimento).")

        bold = ("[b]Negrito:[/b]\n\nPara ativar o negrito basta dizer"
        " [b]\"ativar negrito\"[/b] que tudo o que for dito em seguida"
        " ficará em negrito. Para desativá-lo basta dizer "
        "[b]\"desativar negrito\"[/b] e continuar falando normalmente.")

        italic = ("[b]Itálico:[/b]\n\nPara ativar o itálico basta dizer"
        " [b]\"ativar itálico\"[/b] que tudo o que for dito em seguida"
        " ficará em itálico. Para desativá-lo basta dizer "
        "[b]\"desativar itálico\"[/b] e continuar falando normalmente.")

        underline = ("[b]Sublinhado:[/b]\n\nPara ativar o sublinhado basta dizer"
        " [b]\"ativar sublinhado\"[/b] que tudo o que for dito em seguida"
        " ficará em sublinhado. Para desativá-lo basta dizer "
        "[b]\"desativar sublinhado\"[/b] e continuar falando normalmente.")

        paragraph = ("[b]Novo parágrafo:[/b]\n\nPara adicionar um novo parágrafo"
        " basta dizer [b]\"novo parágrafo\"[/b] que o sistema automaticamente"
        " colocará tudo o que for dito em seguida em um novo parágrafo")

        delParagraph = ("[b]Apagar parágrafo:[/b]\n\nPara apagar o último parágrafo"
        " do seu texto basta dizer [b]\"apagar parágrafo\"[/b].")

        correcPhease = ("[b]Corrigir frase:[/b]\n\nCaso você precise corrigir"
        "a última frase dita, basta dizer [b]\"corrigir por...\"[/b]."
        " seguido da nova frase que o sistema deve substituir.")

        self.ids.textBoxIntructions.add_widget(LabelBox(
            text=title,
            underline=True,
            font_name='Roboto-Bold',
            font_size=20,
            color=[0.4, 0.4, 0.4, 1],
            halign='center'
        ))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=intro, halign='justify'))
        self.ids.textBoxIntructions.add_widget(LabelBox(
            text=subTitle,
            underline=True,
            font_name='Roboto-Bold',
            font_size=20,
            color=[0.4, 0.4, 0.4, 1],
            halign='center'
        ))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=bold, halign='justify'))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=italic, halign='justify'))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=underline, halign='justify'))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=paragraph, halign='justify'))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=delParagraph, halign='justify'))
        self.ids.textBoxIntructions.add_widget(
            LabelBox(text=correcPhease, halign='justify'))


class UserInput(TextInput):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)


class LabelBox(Label):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    def on_size(self, *args) -> None:
        if self.width > 1024:
            self.text_size = (self.width / 1.6, None)
        else:
            self.text_size = (self.width - sp(50), None)

    def on_texture_size(self, *args) -> None:
        self.size = self.texture_size
        self.height += sp(20)


class Yumi(App):
    def build(self):
        with m as source:
            # Ajusta o ruído
            r.adjust_for_ambient_noise(source)
            print("Definindo limite mínimo de percepção para {}".format(
                r.energy_threshold))
        return Manager()


if __name__ == '__main__':
    if hasattr(sys, '_MEIPASS'):
        resource_add_path(os.path.join(sys._MEIPASS))
    Yumi().run()
